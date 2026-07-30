from __future__ import annotations

import asyncio
import os
import re
import shutil
import subprocess
import sys
import tempfile
import traceback
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import edge_tts
import imageio_ffmpeg
from docx import Document
from mutagen.id3 import ID3, ID3NoHeaderError, SYLT, USLT
from PySide6.QtCore import Qt, QThread, Signal
from PySide6.QtGui import QAction
from PySide6.QtWidgets import (
    QApplication,
    QCheckBox,
    QComboBox,
    QFileDialog,
    QGridLayout,
    QGroupBox,
    QHeaderView,
    QHBoxLayout,
    QLabel,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMessageBox,
    QPlainTextEdit,
    QProgressBar,
    QPushButton,
    QSlider,
    QStatusBar,
    QTableWidget,
    QTableWidgetItem,
    QVBoxLayout,
    QWidget,
    QLineEdit,
)


APP_TITLE = "Edge TTS Converter"
SUPPORTED_EXTENSIONS = {".txt", ".docx", ".docm"}
MODE_AUTO = "auto"
MODE_READING = "reading"
MODE_DIALOGUE = "dialogue"
NARRATOR_KEY = "__narrator__"
FOLLOW_ALONG_PAUSE_MULTIPLIER = 1.2

PRESET_VOICES = [
    "en-US-AvaNeural",
    "en-US-JennyNeural",
    "en-US-GuyNeural",
    "en-US-AriaNeural",
    "en-GB-RyanNeural",
    "en-GB-SoniaNeural",
    "en-GB-LibbyNeural",
    "zh-CN-XiaoxiaoNeural",
    "zh-CN-YunxiNeural",
    "zh-CN-YunjianNeural",
    "zh-CN-XiaoyiNeural",
    "de-DE-KatjaNeural",
    "de-DE-ConradNeural",
    "fr-FR-DeniseNeural",
    "es-ES-ElviraNeural",
]

ENGLISH_DIALOGUE_VOICES = [
    "en-US-GuyNeural",
    "en-US-JennyNeural",
    "en-GB-RyanNeural",
    "en-GB-SoniaNeural",
    "en-US-AriaNeural",
    "en-GB-LibbyNeural",
]

CHINESE_DIALOGUE_VOICES = [
    "zh-CN-YunxiNeural",
    "zh-CN-XiaoxiaoNeural",
    "zh-CN-YunjianNeural",
    "zh-CN-XiaoyiNeural",
]

SPEAKER_TAG_PATTERN = re.compile(
    r"^\s*\[speaker(?:\s*:\s*(?P<label>[^\]\r\n]{1,40}))?\]\s*:\s*(?P<text>.+?)\s*$",
    re.IGNORECASE,
)
NARRATION_TAG_PATTERN = re.compile(
    r"^\s*\[narration\]\s*:\s*(?P<text>.+?)\s*$",
    re.IGNORECASE,
)
SENTENCE_ENDINGS = ".!?。！？"
SENTENCE_CLOSERS = "\"'”’）)]】"
NON_ENDING_ABBREVIATIONS = {
    "dr.",
    "jr.",
    "mr.",
    "mrs.",
    "ms.",
    "prof.",
    "sr.",
    "st.",
    "vs.",
}


@dataclass(frozen=True)
class ConvertJob:
    source: Path
    output: Path
    mode: str
    reading_voice: str
    narrator_voice: str
    role_voices: tuple[tuple[str, str], ...]
    rate: int
    dialogue_pause_ms: int
    follow_along: bool
    generate_lyrics: bool


@dataclass(frozen=True)
class SpeechSegment:
    speaker: str
    speaker_key: str
    text: str
    is_narrator: bool


@dataclass(frozen=True)
class LyricCue:
    start_ms: int
    end_ms: int
    text: str


@dataclass(frozen=True)
class RenderedSegment:
    path: Path
    duration_ms: int
    cues: tuple[LyricCue, ...]


@dataclass(frozen=True)
class TextAnalysis:
    is_dialogue: bool
    segments: tuple[SpeechSegment, ...]
    speakers: tuple[str, ...]
    tagged_turns: int
    ignored_lines: int


def normalize_speaker_key(label: str) -> str:
    return " ".join(label.split()).casefold()


def match_tagged_line(line: str) -> tuple[str, str, bool] | None:
    narration_match = NARRATION_TAG_PATTERN.match(line)
    if narration_match:
        return "Narrator", narration_match.group("text").strip(), True

    speaker_match = SPEAKER_TAG_PATTERN.match(line)
    if not speaker_match:
        return None
    label = " ".join((speaker_match.group("label") or "Speaker").split())
    return label, speaker_match.group("text").strip(), False


def append_segment(
    segments: list[SpeechSegment],
    speaker: str,
    speaker_key: str,
    text: str,
    is_narrator: bool,
) -> None:
    text = clean_text(text)
    if not text:
        return
    segments.append(
        SpeechSegment(
            speaker=speaker,
            speaker_key=speaker_key,
            text=text,
            is_narrator=is_narrator,
        )
    )


def analyze_text(text: str, mode: str = MODE_AUTO) -> TextAnalysis:
    if mode == MODE_READING:
        cleaned = clean_text(text)
        segments = (
            SpeechSegment("Narrator", NARRATOR_KEY, cleaned, True),
        ) if cleaned else ()
        return TextAnalysis(False, segments, (), 0, 0)

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    segments_list: list[SpeechSegment] = []
    speaker_names: dict[str, str] = {}
    ignored_lines = 0

    for line in lines:
        if not line.strip():
            continue
        matched = match_tagged_line(line)
        if not matched:
            ignored_lines += 1
            continue
        label, spoken_text, is_narrator = matched
        if is_narrator:
            key = NARRATOR_KEY
            display_name = "Narrator"
        else:
            key = normalize_speaker_key(label)
            speaker_names.setdefault(key, label)
            display_name = speaker_names[key]
        append_segment(segments_list, display_name, key, spoken_text, is_narrator)

    tagged_turns = len(segments_list)
    is_dialogue = tagged_turns > 0 or mode == MODE_DIALOGUE

    if not is_dialogue:
        cleaned = clean_text(text)
        segments = (
            SpeechSegment("Narrator", NARRATOR_KEY, cleaned, True),
        ) if cleaned else ()
        return TextAnalysis(False, segments, (), 0, 0)

    ordered_speakers: list[str] = []
    seen_speakers: set[str] = set()
    for segment in segments_list:
        if segment.is_narrator or segment.speaker_key in seen_speakers:
            continue
        seen_speakers.add(segment.speaker_key)
        ordered_speakers.append(segment.speaker)

    return TextAnalysis(
        True,
        tuple(segments_list),
        tuple(ordered_speakers),
        tagged_turns,
        ignored_lines,
    )


def split_sentences(text: str) -> list[str]:
    cleaned = clean_text(text)
    if not cleaned:
        return []

    sentences: list[str] = []
    start = 0
    index = 0
    while index < len(cleaned):
        ending = cleaned[index]
        if ending not in SENTENCE_ENDINGS:
            index += 1
            continue

        boundary_end = index + 1
        while boundary_end < len(cleaned) and cleaned[boundary_end] in SENTENCE_ENDINGS:
            boundary_end += 1
        while boundary_end < len(cleaned) and cleaned[boundary_end] in SENTENCE_CLOSERS:
            boundary_end += 1

        is_ascii_ending = ending in ".!?"
        has_following_text = boundary_end < len(cleaned)
        if is_ascii_ending and has_following_text and not cleaned[boundary_end].isspace():
            index = boundary_end
            continue

        fragment = cleaned[start : index + 1]
        word_match = re.search(r"([A-Za-z.]+)$", fragment)
        word = word_match.group(1).casefold() if word_match else ""
        is_initialism = bool(re.fullmatch(r"(?:[a-z]\.){2,}", word))
        is_abbreviation = ending == "." and has_following_text and (
            word in NON_ENDING_ABBREVIATIONS or is_initialism
        )
        if is_abbreviation:
            index = boundary_end
            continue

        sentence = clean_text(cleaned[start:boundary_end])
        if sentence:
            sentences.append(sentence)
        start = boundary_end
        while start < len(cleaned) and cleaned[start].isspace():
            start += 1
        index = start

    remainder = clean_text(cleaned[start:])
    if remainder:
        sentences.append(remainder)
    return sentences


def expand_segments_to_sentences(
    segments: tuple[SpeechSegment, ...],
) -> tuple[SpeechSegment, ...]:
    expanded: list[SpeechSegment] = []
    for segment in segments:
        for sentence in split_sentences(segment.text):
            expanded.append(
                SpeechSegment(
                    speaker=segment.speaker,
                    speaker_key=segment.speaker_key,
                    text=sentence,
                    is_narrator=segment.is_narrator,
                )
            )
    return tuple(expanded)


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    lines = [" ".join(line.split()) for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def read_txt(path: Path) -> str:
    encodings = ["utf-8-sig", "utf-8", "utf-16", "gb18030", "big5", "cp1252"]
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_input_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return read_txt(path)
    if suffix in {".docx", ".docm"}:
        return read_docx(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def rate_to_edge_value(rate: int) -> str:
    sign = "+" if rate >= 0 else ""
    return f"{sign}{rate}%"


def safe_output_path(
    output_dir: Path,
    source: Path,
    reserved_paths: set[Path] | None = None,
) -> Path:
    base = re.sub(r'[<>:"/\\|?*]+', "_", source.stem).strip() or "tts-output"
    candidate = output_dir / f"{base}.mp3"
    index = 2
    reserved_paths = reserved_paths if reserved_paths is not None else set()
    while candidate.exists() or candidate.resolve() in reserved_paths:
        candidate = output_dir / f"{base}_{index}.mp3"
        index += 1
    reserved_paths.add(candidate.resolve())
    return candidate


async def synthesize_to_mp3(
    text: str,
    voice: str,
    rate: int,
    output: Path,
) -> tuple[LyricCue, ...]:
    communicate = edge_tts.Communicate(
        text=text,
        voice=voice,
        rate=rate_to_edge_value(rate),
        boundary="SentenceBoundary",
    )
    cues: list[LyricCue] = []
    with output.open("wb") as audio_file:
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_file.write(chunk["data"])
            elif chunk["type"] == "SentenceBoundary":
                start_ms = round(chunk["offset"] / 10_000)
                duration_ms = round(chunk["duration"] / 10_000)
                cues.append(
                    LyricCue(
                        start_ms=start_ms,
                        end_ms=start_ms + duration_ms,
                        text=clean_text(chunk["text"]),
                    )
                )
    return tuple(cues)


def run_ffmpeg(arguments: list[str]) -> None:
    creation_flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    result = subprocess.run(
        [imageio_ffmpeg.get_ffmpeg_exe(), *arguments],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        creationflags=creation_flags,
        check=False,
    )
    if result.returncode != 0:
        details = result.stderr.strip().splitlines()
        message = details[-1] if details else "Unknown FFmpeg error"
        raise RuntimeError(f"Could not combine dialogue audio: {message}")


def probe_audio_duration_ms(path: Path) -> int:
    creation_flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    result = subprocess.run(
        [imageio_ffmpeg.get_ffmpeg_exe(), "-hide_banner", "-i", str(path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        creationflags=creation_flags,
        check=False,
    )
    match = re.search(r"Duration:\s*(\d+):(\d+):(\d+(?:\.\d+)?)", result.stderr)
    if not match:
        raise RuntimeError(f"Could not measure sentence duration: {path.name}")
    hours, minutes, seconds = match.groups()
    total_seconds = int(hours) * 3600 + int(minutes) * 60 + float(seconds)
    return max(1, round(total_seconds * 1000))


def create_silence_mp3(output: Path, duration_ms: int) -> None:
    run_ffmpeg(
        [
            "-hide_banner",
            "-loglevel",
            "error",
            "-y",
            "-f",
            "lavfi",
            "-i",
            "anullsrc=r=24000:cl=mono",
            "-t",
            f"{duration_ms / 1000:.3f}",
            "-ar",
            "24000",
            "-ac",
            "1",
            "-c:a",
            "libmp3lame",
            "-b:a",
            "48k",
            str(output),
        ]
    )


def encode_concat_list(sequence: list[Path], list_name: str, output: Path) -> None:
    work_dir = sequence[0].parent
    concat_file = work_dir / list_name
    concat_file.write_text(
        "\n".join(f"file '{path.name}'" for path in sequence),
        encoding="utf-8",
    )
    run_ffmpeg(
        [
            "-hide_banner",
            "-loglevel",
            "error",
            "-y",
            "-f",
            "concat",
            "-safe",
            "0",
            "-i",
            str(concat_file),
            "-ar",
            "24000",
            "-ac",
            "1",
            "-c:a",
            "libmp3lame",
            "-b:a",
            "48k",
            str(output),
        ]
    )


def dialogue_transition_pause_ms(
    index: int,
    segments: tuple[SpeechSegment, ...],
    pause_ms: int,
) -> int:
    if pause_ms <= 0 or index >= len(segments) - 1:
        return 0
    uses_narrator = segments[index].is_narrator or segments[index + 1].is_narrator
    return min(2500, pause_ms + 250) if uses_narrator else pause_ms


def combine_dialogue_mp3(
    segment_files: list[Path],
    segments: tuple[SpeechSegment, ...],
    pause_ms: int,
    output: Path,
) -> list[int]:
    if len(segment_files) == 1:
        shutil.copyfile(segment_files[0], output)
        return [0]

    work_dir = segment_files[0].parent
    standard_pause = work_dir / "pause.mp3"
    narrator_pause = work_dir / "narrator-pause.mp3"
    if pause_ms > 0:
        create_silence_mp3(standard_pause, pause_ms)
        create_silence_mp3(narrator_pause, min(2500, pause_ms + 250))
    standard_duration_ms = probe_audio_duration_ms(standard_pause) if standard_pause.exists() else 0
    narrator_duration_ms = probe_audio_duration_ms(narrator_pause) if narrator_pause.exists() else 0

    sequence: list[Path] = []
    applied_pauses_ms: list[int] = []
    for index, segment_file in enumerate(segment_files):
        sequence.append(segment_file)
        transition_pause_ms = dialogue_transition_pause_ms(index, segments, pause_ms)
        if transition_pause_ms <= 0:
            applied_pauses_ms.append(0)
            continue
        pause_file = narrator_pause if transition_pause_ms > pause_ms else standard_pause
        sequence.append(pause_file)
        applied_pauses_ms.append(
            narrator_duration_ms if pause_file == narrator_pause else standard_duration_ms
        )

    encode_concat_list(sequence, "dialogue-concat.txt", output)
    return applied_pauses_ms


def combine_follow_along_mp3(
    segment_files: list[Path],
    pause_durations_ms: list[int],
    output: Path,
) -> list[int]:
    if len(segment_files) != len(pause_durations_ms):
        raise ValueError("Sentence audio and pause counts do not match")
    sequence: list[Path] = []
    applied_pauses_ms: list[int] = []
    work_dir = segment_files[0].parent
    for index, (segment_file, duration_ms) in enumerate(
        zip(segment_files, pause_durations_ms),
        start=1,
    ):
        pause_file = work_dir / f"follow-pause-{index:04d}.mp3"
        create_silence_mp3(pause_file, duration_ms)
        applied_pauses_ms.append(probe_audio_duration_ms(pause_file))
        sequence.extend([segment_file, pause_file])
    encode_concat_list(sequence, "follow-along-concat.txt", output)
    return applied_pauses_ms


async def synthesize_dialogue_segments(
    segments: tuple[SpeechSegment, ...],
    role_voices: dict[str, str],
    narrator_voice: str,
    fallback_voice: str,
    rate: int,
    work_dir: Path,
    progress_callback: Callable[[int, int, SpeechSegment], None] | None = None,
) -> list[RenderedSegment]:
    rendered_segments: list[RenderedSegment] = []
    total = len(segments)
    for index, segment in enumerate(segments, start=1):
        voice = narrator_voice if segment.is_narrator else role_voices.get(
            segment.speaker_key,
            fallback_voice,
        )
        segment_file = work_dir / f"segment-{index:04d}.mp3"
        if progress_callback:
            progress_callback(index, total, segment)
        cues = await synthesize_to_mp3(segment.text, voice, rate, segment_file)
        duration_ms = probe_audio_duration_ms(segment_file)
        if not cues:
            cues = (LyricCue(0, duration_ms, segment.text),)
        rendered_segments.append(
            RenderedSegment(
                path=segment_file,
                duration_ms=duration_ms,
                cues=cues,
            )
        )
    return rendered_segments


def displayed_lyric_text(
    segment: SpeechSegment,
    text: str,
    is_dialogue: bool,
) -> str:
    if not is_dialogue:
        return text
    if segment.is_narrator:
        return f"Narration: {text}"
    if segment.speaker == "Speaker":
        return text
    return f"{segment.speaker}: {text}"


def build_dialogue_lyric_cues(
    rendered_segments: list[RenderedSegment],
    segments: tuple[SpeechSegment, ...],
    applied_pauses_ms: list[int],
) -> tuple[LyricCue, ...]:
    if not (
        len(rendered_segments) == len(segments) == len(applied_pauses_ms)
    ):
        raise ValueError("Dialogue segments and pause counts do not match")
    cues: list[LyricCue] = []
    timeline_ms = 0
    for index, (rendered, segment) in enumerate(zip(rendered_segments, segments)):
        for cue in rendered.cues:
            cues.append(
                LyricCue(
                    start_ms=timeline_ms + cue.start_ms,
                    end_ms=timeline_ms + cue.end_ms,
                    text=displayed_lyric_text(segment, cue.text, True),
                )
            )
        timeline_ms += rendered.duration_ms
        timeline_ms += applied_pauses_ms[index]
    return tuple(cues)


def build_follow_along_lyric_cues(
    rendered_segments: list[RenderedSegment],
    segments: tuple[SpeechSegment, ...],
    is_dialogue: bool,
    applied_pauses_ms: list[int],
) -> tuple[LyricCue, ...]:
    if not (
        len(rendered_segments) == len(segments) == len(applied_pauses_ms)
    ):
        raise ValueError("Follow-along segments and pause counts do not match")
    cues: list[LyricCue] = []
    timeline_ms = 0
    for index, (rendered, segment) in enumerate(zip(rendered_segments, segments)):
        for cue in rendered.cues:
            cues.append(
                LyricCue(
                    start_ms=timeline_ms + cue.start_ms,
                    end_ms=timeline_ms + cue.end_ms,
                    text=displayed_lyric_text(segment, cue.text, is_dialogue),
                )
            )
        timeline_ms += rendered.duration_ms
        timeline_ms += applied_pauses_ms[index]
    return tuple(cues)


def format_lrc_timestamp(milliseconds: int) -> str:
    centiseconds = max(0, milliseconds) // 10
    minutes, remainder = divmod(centiseconds, 6000)
    seconds, fraction = divmod(remainder, 100)
    return f"{minutes:02d}:{seconds:02d}.{fraction:02d}"


def lyric_language(cues: tuple[LyricCue, ...]) -> str:
    text = " ".join(cue.text for cue in cues)
    cjk_count = len(re.findall(r"[\u3400-\u9fff]", text))
    latin_count = len(re.findall(r"[A-Za-z]", text))
    return "zho" if cjk_count > latin_count else "eng"


def write_synchronized_lyrics(
    mp3_path: Path,
    cues: tuple[LyricCue, ...],
) -> Path:
    if not cues:
        raise ValueError("No synchronized lyric cues were generated")

    lrc_path = mp3_path.with_suffix(".lrc")
    lrc_lines = [
        f"[ti:{mp3_path.stem}]",
        "[by:Edge TTS Converter]",
        *(
            f"[{format_lrc_timestamp(cue.start_ms)}]{cue.text}"
            for cue in cues
        ),
    ]
    lrc_path.write_text("\n".join(lrc_lines) + "\n", encoding="utf-8-sig")

    try:
        tags = ID3(mp3_path)
    except ID3NoHeaderError:
        tags = ID3()
    tags.delall("USLT")
    tags.delall("SYLT")
    language = lyric_language(cues)
    tags.add(
        USLT(
            encoding=1,
            lang=language,
            desc="Lyrics",
            text="\n".join(cue.text for cue in cues),
        )
    )
    tags.add(
        SYLT(
            encoding=1,
            lang=language,
            format=2,
            type=1,
            desc="Synchronized lyrics",
            text=[(cue.text, cue.start_ms) for cue in cues],
        )
    )
    tags.save(mp3_path, v2_version=3)
    return lrc_path


def synthesize_dialogue_to_mp3(
    analysis: TextAnalysis,
    role_voices: dict[str, str],
    narrator_voice: str,
    fallback_voice: str,
    rate: int,
    pause_ms: int,
    output: Path,
    progress_callback: Callable[[int, int, SpeechSegment], None] | None = None,
) -> tuple[LyricCue, ...]:
    with tempfile.TemporaryDirectory(prefix="edge-tts-dialogue-") as tmp:
        work_dir = Path(tmp)
        rendered_segments = asyncio.run(
            synthesize_dialogue_segments(
                analysis.segments,
                role_voices,
                narrator_voice,
                fallback_voice,
                rate,
                work_dir,
                progress_callback,
            )
        )
        segment_files = [rendered.path for rendered in rendered_segments]
        combined_output = work_dir / "combined.mp3"
        applied_pauses_ms = combine_dialogue_mp3(
            segment_files,
            analysis.segments,
            pause_ms,
            combined_output,
        )
        shutil.copyfile(combined_output, output)
        return build_dialogue_lyric_cues(
            rendered_segments,
            analysis.segments,
            applied_pauses_ms,
        )


def synthesize_follow_along_to_mp3(
    segments: tuple[SpeechSegment, ...],
    role_voices: dict[str, str],
    narrator_voice: str,
    fallback_voice: str,
    rate: int,
    is_dialogue: bool,
    output: Path,
    progress_callback: Callable[[int, int, SpeechSegment], None] | None = None,
) -> tuple[LyricCue, ...]:
    sentence_segments = expand_segments_to_sentences(segments)
    if not sentence_segments:
        raise ValueError("No speakable sentences found")

    with tempfile.TemporaryDirectory(prefix="edge-tts-follow-along-") as tmp:
        work_dir = Path(tmp)
        rendered_segments = asyncio.run(
            synthesize_dialogue_segments(
                sentence_segments,
                role_voices,
                narrator_voice,
                fallback_voice,
                rate,
                work_dir,
                progress_callback,
            )
        )
        segment_files = [rendered.path for rendered in rendered_segments]
        pause_durations_ms = [
            round(rendered.duration_ms * FOLLOW_ALONG_PAUSE_MULTIPLIER)
            for rendered in rendered_segments
        ]
        combined_output = work_dir / "combined.mp3"
        applied_pauses_ms = combine_follow_along_mp3(
            segment_files,
            pause_durations_ms,
            combined_output,
        )
        shutil.copyfile(combined_output, output)
        return build_follow_along_lyric_cues(
            rendered_segments,
            sentence_segments,
            is_dialogue,
            applied_pauses_ms,
        )


def run_self_test() -> int:
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        source = tmpdir / "self-test.txt"
        source.write_text(
            "[speaker: Ava]: This is a short dialogue self test.\n"
            "[speaker: Ryan]: The audio tools are working.",
            encoding="utf-8",
        )
        output = tmpdir / "self-test.mp3"
        analysis = analyze_text(read_input_file(source), MODE_AUTO)
        if not analysis.is_dialogue:
            return 1
        cues = synthesize_follow_along_to_mp3(
            segments=analysis.segments,
            role_voices={"ava": "en-US-AvaNeural", "ryan": "en-GB-RyanNeural"},
            narrator_voice="en-US-AvaNeural",
            fallback_voice="en-US-AvaNeural",
            rate=-8,
            is_dialogue=True,
            output=output,
        )
        lrc_path = write_synchronized_lyrics(output, cues)
        if not output.exists() or output.stat().st_size == 0 or not lrc_path.exists():
            return 1
        tags = ID3(output)
        if not tags.getall("SYLT") or not tags.getall("USLT"):
            return 1
    return 0


class ConvertWorker(QThread):
    progress = Signal(int, int)
    status = Signal(str)
    log = Signal(str)
    finished_summary = Signal(int, int)

    def __init__(self, jobs: list[ConvertJob]) -> None:
        super().__init__()
        self.jobs = jobs

    def run(self) -> None:
        completed = 0
        failed = 0
        total = len(self.jobs)

        for index, job in enumerate(self.jobs, start=1):
            try:
                self.status.emit(f"Reading {job.source.name}")
                source_text = read_input_file(job.source)
                if not clean_text(source_text):
                    raise ValueError("No readable text found")

                analysis = analyze_text(source_text, job.mode)
                if analysis.is_dialogue:
                    if not analysis.segments:
                        raise ValueError(
                            "No valid [speaker]: or [narration]: lines were found"
                        )
                    self.log.emit(
                        f"Dialogue detected: {job.source.name} - "
                        f"{len(analysis.speakers)} speaker(s), {len(analysis.segments)} turn(s)"
                    )
                    if analysis.ignored_lines:
                        self.log.emit(
                            f"Skipped {analysis.ignored_lines} unmarked line(s) in "
                            f"{job.source.name}"
                        )

                role_voices = dict(job.role_voices)
                lyric_cues: tuple[LyricCue, ...] = ()
                if job.follow_along:
                    sentence_segments = expand_segments_to_sentences(analysis.segments)

                    def report_sentence(
                        done: int,
                        sentence_total: int,
                        segment: SpeechSegment,
                    ) -> None:
                        self.status.emit(
                            f"{job.source.name}: sentence {done}/{sentence_total} "
                            f"({segment.speaker})"
                        )

                    self.log.emit(
                        f"Follow-along mode: {job.source.name} - "
                        f"{len(sentence_segments)} sentence(s), 1.2x pauses"
                    )
                    lyric_cues = synthesize_follow_along_to_mp3(
                        segments=analysis.segments,
                        role_voices=role_voices,
                        narrator_voice=(
                            job.narrator_voice if analysis.is_dialogue else job.reading_voice
                        ),
                        fallback_voice=job.reading_voice,
                        rate=job.rate,
                        is_dialogue=analysis.is_dialogue,
                        output=job.output,
                        progress_callback=report_sentence,
                    )
                elif analysis.is_dialogue:

                    def report_turn(done: int, turn_total: int, segment: SpeechSegment) -> None:
                        self.status.emit(
                            f"{job.source.name}: turn {done}/{turn_total} ({segment.speaker})"
                        )

                    lyric_cues = synthesize_dialogue_to_mp3(
                        analysis=analysis,
                        role_voices=role_voices,
                        narrator_voice=job.narrator_voice,
                        fallback_voice=job.reading_voice,
                        rate=job.rate,
                        pause_ms=job.dialogue_pause_ms,
                        output=job.output,
                        progress_callback=report_turn,
                    )
                else:
                    self.status.emit(f"Converting {job.source.name} as reading")
                    lyric_cues = asyncio.run(
                        synthesize_to_mp3(
                            clean_text(source_text),
                            job.reading_voice,
                            job.rate,
                            job.output,
                        )
                    )
                if job.generate_lyrics:
                    if not lyric_cues:
                        duration_ms = probe_audio_duration_ms(job.output)
                        lyric_cues = (
                            LyricCue(
                                0,
                                duration_ms,
                                clean_text(source_text),
                            ),
                        )
                    lrc_path = write_synchronized_lyrics(job.output, lyric_cues)
                    self.log.emit(
                        f"Lyrics: {lrc_path.name} + embedded synchronized lyrics"
                    )
                completed += 1
                self.log.emit(f"Done: {job.source.name} -> {job.output.name}")
            except Exception as exc:
                failed += 1
                self.log.emit(f"Failed: {job.source.name} - {exc}")
                self.log.emit(traceback.format_exc())
            finally:
                self.progress.emit(index, total)

        self.finished_summary.emit(completed, failed)


class VoiceRefreshWorker(QThread):
    voices_loaded = Signal(list)
    failed = Signal(str)

    def run(self) -> None:
        try:
            voices = asyncio.run(edge_tts.list_voices())
            names = sorted({voice["ShortName"] for voice in voices if "ShortName" in voice})
            self.voices_loaded.emit(names)
        except Exception as exc:
            self.failed.emit(f"Could not refresh voices: {exc}")


class EdgeTTSConverterWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle(APP_TITLE)
        self.resize(1020, 800)
        self.selected_files: list[Path] = []
        self.available_voices = list(PRESET_VOICES)
        self.role_voice_choices: dict[str, str] = {}
        self.detected_speaker_samples: dict[str, tuple[str, str]] = {}
        self.convert_worker: ConvertWorker | None = None
        self.voice_worker: VoiceRefreshWorker | None = None

        self._build_ui()
        self._build_menu()
        self._update_progress_label(0, 0)
        self.statusBar().showMessage("Ready")

    def _build_menu(self) -> None:
        help_menu = self.menuBar().addMenu("Help")
        open_prompt = QAction("Open Dialogue Prompt", self)
        open_prompt.triggered.connect(self.open_dialogue_prompt)
        help_menu.addAction(open_prompt)
        about = QAction("About", self)
        about.triggered.connect(self.show_about)
        help_menu.addAction(about)

    def _build_ui(self) -> None:
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)

        file_bar = QHBoxLayout()
        self.add_button = QPushButton("Add Files")
        self.remove_button = QPushButton("Remove Selected")
        self.clear_button = QPushButton("Clear")
        file_bar.addWidget(self.add_button)
        file_bar.addWidget(self.remove_button)
        file_bar.addWidget(self.clear_button)
        file_bar.addStretch(1)
        main_layout.addLayout(file_bar)

        settings_group = QGroupBox("Settings")
        settings_layout = QGridLayout(settings_group)
        main_layout.addWidget(settings_group)

        self.output_edit = QLineEdit(str(Path.home() / "Desktop" / "EdgeTTS_MP3"))
        self.output_button = QPushButton("Browse")
        settings_layout.addWidget(QLabel("Output folder"), 0, 0)
        settings_layout.addWidget(self.output_edit, 0, 1)
        settings_layout.addWidget(self.output_button, 0, 2)

        self.mode_combo = QComboBox()
        self.mode_combo.addItem("Auto detect", MODE_AUTO)
        self.mode_combo.addItem("Reading", MODE_READING)
        self.mode_combo.addItem("Dialogue", MODE_DIALOGUE)
        settings_layout.addWidget(QLabel("Text mode"), 1, 0)
        settings_layout.addWidget(self.mode_combo, 1, 1, 1, 2)

        self.voice_combo = QComboBox()
        self.voice_combo.setEditable(True)
        self.voice_combo.addItems(self.available_voices)
        self.refresh_voices_button = QPushButton("Refresh Voices")
        settings_layout.addWidget(QLabel("Reading voice"), 2, 0)
        settings_layout.addWidget(self.voice_combo, 2, 1)
        settings_layout.addWidget(self.refresh_voices_button, 2, 2)

        self.narrator_voice_combo = QComboBox()
        self.narrator_voice_combo.setEditable(True)
        self.narrator_voice_combo.addItems(self.available_voices)
        self.narrator_voice_combo.setCurrentText("en-US-AvaNeural")
        settings_layout.addWidget(QLabel("Narrator voice"), 3, 0)
        settings_layout.addWidget(self.narrator_voice_combo, 3, 1, 1, 2)

        self.speed_slider = QSlider(Qt.Orientation.Horizontal)
        self.speed_slider.setRange(-50, 50)
        self.speed_slider.setValue(0)
        self.speed_label = QLabel(rate_to_edge_value(0))
        self.speed_label.setMinimumWidth(60)
        settings_layout.addWidget(QLabel("Speed"), 4, 0)
        settings_layout.addWidget(self.speed_slider, 4, 1)
        settings_layout.addWidget(self.speed_label, 4, 2)

        self.pause_slider = QSlider(Qt.Orientation.Horizontal)
        self.pause_slider.setRange(0, 2000)
        self.pause_slider.setSingleStep(50)
        self.pause_slider.setPageStep(100)
        self.pause_slider.setValue(550)
        self.pause_label = QLabel("550 ms")
        self.pause_label.setMinimumWidth(60)
        self.turn_pause_title = QLabel("Turn pause")
        settings_layout.addWidget(self.turn_pause_title, 5, 0)
        settings_layout.addWidget(self.pause_slider, 5, 1)
        settings_layout.addWidget(self.pause_label, 5, 2)

        self.follow_along_checkbox = QCheckBox("Follow-along mode (1.2x sentence pause)")
        settings_layout.addWidget(self.follow_along_checkbox, 6, 1, 1, 2)

        self.lyrics_checkbox = QCheckBox("Synchronized lyrics (.lrc + embedded)")
        self.lyrics_checkbox.setChecked(True)
        settings_layout.addWidget(self.lyrics_checkbox, 7, 1, 1, 2)

        self.file_list = QListWidget()
        self.file_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        main_layout.addWidget(self.file_list, stretch=2)

        dialogue_group = QGroupBox("Dialogue Voices")
        dialogue_layout = QVBoxLayout(dialogue_group)
        self.role_table = QTableWidget(0, 2)
        self.role_table.setHorizontalHeaderLabels(["Speaker", "Voice"])
        self.role_table.verticalHeader().setVisible(False)
        self.role_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.role_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.role_table.setMinimumHeight(135)
        dialogue_layout.addWidget(self.role_table)
        main_layout.addWidget(dialogue_group, stretch=1)

        action_bar = QHBoxLayout()
        self.start_button = QPushButton("Start Convert")
        self.open_output_button = QPushButton("Open Output Folder")
        action_bar.addWidget(self.start_button)
        action_bar.addWidget(self.open_output_button)
        action_bar.addStretch(1)
        main_layout.addLayout(action_bar)

        log_group = QGroupBox("Log")
        log_layout = QVBoxLayout(log_group)
        self.log_text = QPlainTextEdit()
        self.log_text.setReadOnly(True)
        log_layout.addWidget(self.log_text)
        main_layout.addWidget(log_group, stretch=1)

        status = QStatusBar()
        self.setStatusBar(status)
        self.progress_label = QLabel()
        self.progress_bar = QProgressBar()
        self.progress_bar.setMinimumWidth(260)
        status.addPermanentWidget(self.progress_label)
        status.addPermanentWidget(self.progress_bar)

        self.add_button.clicked.connect(self.add_files)
        self.remove_button.clicked.connect(self.remove_selected)
        self.clear_button.clicked.connect(self.clear_files)
        self.output_button.clicked.connect(self.choose_output_dir)
        self.refresh_voices_button.clicked.connect(self.refresh_voices)
        self.mode_combo.currentIndexChanged.connect(self.refresh_detected_speakers)
        self.speed_slider.valueChanged.connect(self.update_speed_label)
        self.pause_slider.valueChanged.connect(self.update_pause_label)
        self.follow_along_checkbox.toggled.connect(self.update_follow_along_state)
        self.start_button.clicked.connect(self.start_conversion)
        self.open_output_button.clicked.connect(self.open_output_folder)

    def show_about(self) -> None:
        QMessageBox.information(
            self,
            APP_TITLE,
            "Convert TXT and Word DOCX files to MP3 with Microsoft Edge TTS.\n"
            "Dialogue lines must use [speaker]: or [speaker: Name]: tags.\n"
            "Narration lines must use [narration]: tags.\n"
            "Follow-along mode adds a 1.2x sentence-duration pause.\n\n"
            "Synchronized lyrics can be saved as LRC and embedded in MP3.\n\n"
            "Internet access is required. Old .doc files should be saved as .docx first.",
        )

    def open_dialogue_prompt(self) -> None:
        app_dir = Path(sys.executable).parent if getattr(sys, "frozen", False) else Path(__file__).parent
        prompt_path = app_dir / "dialogue_prompt.txt"
        if prompt_path.exists():
            os.startfile(str(prompt_path))
            return
        QMessageBox.warning(self, APP_TITLE, f"Prompt file not found: {prompt_path}")

    def add_files(self) -> None:
        paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Select text or Word files",
            "",
            "Text and Word files (*.txt *.docx *.docm);;Text files (*.txt);;Word files (*.docx *.docm);;All files (*.*)",
        )
        if not paths:
            return

        known = {path.resolve() for path in self.selected_files}
        unsupported: list[str] = []
        added = 0

        for raw_path in paths:
            path = Path(raw_path)
            if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                unsupported.append(path.name)
                continue
            resolved = path.resolve()
            if resolved not in known:
                self.selected_files.append(path)
                known.add(resolved)
                added += 1

        self.refresh_file_list()
        if unsupported:
            self.log(f"Skipped unsupported files: {', '.join(unsupported)}")
        if added:
            self.statusBar().showMessage(f"Added {added} file(s)")

    def remove_selected(self) -> None:
        selected_rows = sorted({index.row() for index in self.file_list.selectedIndexes()}, reverse=True)
        if not selected_rows:
            return
        for row in selected_rows:
            self.selected_files.pop(row)
        self.refresh_file_list()
        self.statusBar().showMessage("Selected file(s) removed")

    def clear_files(self) -> None:
        self.selected_files.clear()
        self.refresh_file_list()
        self.statusBar().showMessage("File list cleared")

    def refresh_file_list(self) -> None:
        self.file_list.clear()
        for path in self.selected_files:
            item = QListWidgetItem(str(path))
            item.setToolTip(str(path))
            self.file_list.addItem(item)
        self.refresh_detected_speakers()
        self._update_progress_label(0, len(self.selected_files))

    def current_mode(self) -> str:
        return self.mode_combo.currentData() or MODE_AUTO

    def default_role_voice(self, index: int, sample_text: str) -> str:
        cjk_count = len(re.findall(r"[\u3400-\u9fff]", sample_text))
        latin_count = len(re.findall(r"[A-Za-z]", sample_text))
        pool = CHINESE_DIALOGUE_VOICES if cjk_count > latin_count else ENGLISH_DIALOGUE_VOICES
        return pool[index % len(pool)]

    def refresh_detected_speakers(self) -> None:
        if not hasattr(self, "role_table"):
            return

        for row in range(self.role_table.rowCount()):
            item = self.role_table.item(row, 0)
            combo = self.role_table.cellWidget(row, 1)
            if item and isinstance(combo, QComboBox):
                key = item.data(Qt.ItemDataRole.UserRole)
                if key:
                    self.role_voice_choices[key] = combo.currentText().strip()

        samples: dict[str, tuple[str, str]] = {}
        mode = self.current_mode()
        if mode != MODE_READING:
            for path in self.selected_files:
                try:
                    analysis = analyze_text(read_input_file(path), mode)
                except Exception as exc:
                    self.log(f"Could not inspect {path.name}: {exc}")
                    continue
                if not analysis.is_dialogue:
                    continue
                for segment in analysis.segments:
                    if segment.is_narrator:
                        continue
                    samples.setdefault(segment.speaker_key, (segment.speaker, segment.text))

        self.detected_speaker_samples = samples
        self.role_table.setRowCount(0)
        for row, (key, (display_name, sample_text)) in enumerate(samples.items()):
            self.role_table.insertRow(row)
            name_item = QTableWidgetItem(display_name)
            name_item.setData(Qt.ItemDataRole.UserRole, key)
            name_item.setFlags(name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            name_item.setToolTip(sample_text[:240])
            self.role_table.setItem(row, 0, name_item)

            voice_combo = QComboBox()
            voice_combo.setEditable(True)
            voice_combo.addItems(self.available_voices)
            selected_voice = self.role_voice_choices.get(key)
            if not selected_voice:
                selected_voice = self.default_role_voice(row, sample_text)
                self.role_voice_choices[key] = selected_voice
            voice_combo.setCurrentText(selected_voice)
            voice_combo.currentTextChanged.connect(
                lambda value, speaker_key=key: self.role_voice_choices.__setitem__(
                    speaker_key,
                    value.strip(),
                )
            )
            self.role_table.setCellWidget(row, 1, voice_combo)

    def replace_combo_voices(self, combo: QComboBox, voices: list[str]) -> None:
        current = combo.currentText().strip()
        combo.blockSignals(True)
        combo.clear()
        combo.addItems(voices)
        combo.setCurrentText(current)
        combo.blockSignals(False)

    def choose_output_dir(self) -> None:
        path = QFileDialog.getExistingDirectory(self, "Choose output folder")
        if path:
            self.output_edit.setText(path)

    def open_output_folder(self) -> None:
        output_dir = Path(self.output_edit.text()).expanduser()
        output_dir.mkdir(parents=True, exist_ok=True)
        os.startfile(str(output_dir))

    def refresh_voices(self) -> None:
        if self.convert_worker and self.convert_worker.isRunning():
            return
        self.statusBar().showMessage("Fetching voice list...")
        self.refresh_voices_button.setEnabled(False)
        self.voice_worker = VoiceRefreshWorker()
        self.voice_worker.voices_loaded.connect(self.on_voices_loaded)
        self.voice_worker.failed.connect(self.on_voice_refresh_failed)
        self.voice_worker.finished.connect(lambda: self.refresh_voices_button.setEnabled(True))
        self.voice_worker.start()

    def on_voices_loaded(self, voices: list[str]) -> None:
        self.available_voices = voices
        self.replace_combo_voices(self.voice_combo, voices)
        self.replace_combo_voices(self.narrator_voice_combo, voices)
        self.refresh_detected_speakers()
        self.log(f"Loaded {len(voices)} voices")
        self.statusBar().showMessage(f"Loaded {len(voices)} voice(s)")

    def on_voice_refresh_failed(self, message: str) -> None:
        self.log(message)
        self.statusBar().showMessage("Voice refresh failed")

    def update_speed_label(self, value: int) -> None:
        self.speed_label.setText(rate_to_edge_value(value))

    def update_pause_label(self, value: int) -> None:
        self.pause_label.setText(f"{value} ms")

    def update_follow_along_state(self, checked: bool) -> None:
        self.turn_pause_title.setEnabled(not checked)
        self.pause_slider.setEnabled(not checked)
        self.pause_label.setEnabled(not checked)

    def start_conversion(self) -> None:
        if self.convert_worker and self.convert_worker.isRunning():
            return
        if not self.selected_files:
            QMessageBox.information(self, APP_TITLE, "Please add one or more TXT or Word files.")
            return

        reading_voice = self.voice_combo.currentText().strip()
        narrator_voice = self.narrator_voice_combo.currentText().strip()
        if not reading_voice or not narrator_voice:
            QMessageBox.information(self, APP_TITLE, "Please choose the reading and narrator voices.")
            return

        self.refresh_detected_speakers()
        role_voices = tuple(
            (key, self.role_voice_choices.get(key, "").strip())
            for key in self.detected_speaker_samples
        )
        missing_roles = [
            display_name
            for key, (display_name, _) in self.detected_speaker_samples.items()
            if not self.role_voice_choices.get(key, "").strip()
        ]
        if missing_roles:
            QMessageBox.information(
                self,
                APP_TITLE,
                f"Please choose a voice for: {', '.join(missing_roles)}",
            )
            return

        output_dir = Path(self.output_edit.text()).expanduser()
        output_dir.mkdir(parents=True, exist_ok=True)
        rate = self.speed_slider.value()
        mode = self.current_mode()
        jobs: list[ConvertJob] = []
        reserved_outputs: set[Path] = set()
        for path in self.selected_files:
            jobs.append(
                ConvertJob(
                    source=path,
                    output=safe_output_path(output_dir, path, reserved_outputs),
                    mode=mode,
                    reading_voice=reading_voice,
                    narrator_voice=narrator_voice,
                    role_voices=role_voices,
                    rate=rate,
                    dialogue_pause_ms=self.pause_slider.value(),
                    follow_along=self.follow_along_checkbox.isChecked(),
                    generate_lyrics=self.lyrics_checkbox.isChecked(),
                )
            )

        self.set_controls_enabled(False)
        self.progress_bar.setMaximum(len(jobs))
        self.progress_bar.setValue(0)
        self._update_progress_label(0, len(jobs))
        self.log(
            f"Starting {len(jobs)} job(s), mode {mode}, speed {rate_to_edge_value(rate)}, "
            f"follow-along {'on' if self.follow_along_checkbox.isChecked() else 'off'}, "
            f"lyrics {'on' if self.lyrics_checkbox.isChecked() else 'off'}"
        )
        self.statusBar().showMessage("Converting...")

        self.convert_worker = ConvertWorker(jobs)
        self.convert_worker.status.connect(self.statusBar().showMessage)
        self.convert_worker.log.connect(self.log)
        self.convert_worker.progress.connect(self.on_progress)
        self.convert_worker.finished_summary.connect(self.on_finished)
        self.convert_worker.start()

    def set_controls_enabled(self, enabled: bool) -> None:
        for widget in [
            self.add_button,
            self.remove_button,
            self.clear_button,
            self.output_button,
            self.refresh_voices_button,
            self.mode_combo,
            self.voice_combo,
            self.narrator_voice_combo,
            self.speed_slider,
            self.follow_along_checkbox,
            self.lyrics_checkbox,
            self.role_table,
            self.start_button,
        ]:
            widget.setEnabled(enabled)
        self.turn_pause_title.setEnabled(enabled and not self.follow_along_checkbox.isChecked())
        self.pause_slider.setEnabled(enabled and not self.follow_along_checkbox.isChecked())
        self.pause_label.setEnabled(enabled and not self.follow_along_checkbox.isChecked())

    def on_progress(self, done: int, total: int) -> None:
        self.progress_bar.setMaximum(total)
        self.progress_bar.setValue(done)
        self._update_progress_label(done, total)

    def on_finished(self, completed: int, failed: int) -> None:
        self.set_controls_enabled(True)
        self.statusBar().showMessage(f"Finished: {completed} done, {failed} failed")
        self.log(f"Finished: {completed} done, {failed} failed")

    def _update_progress_label(self, done: int, total: int) -> None:
        self.progress_label.setText(f"{done} / {total}")

    def log(self, message: str) -> None:
        self.log_text.appendPlainText(message.rstrip())


def main() -> int:
    if "--self-test" in sys.argv:
        return run_self_test()

    app = QApplication(sys.argv)
    window = EdgeTTSConverterWindow()
    window.show()
    return app.exec()


if __name__ == "__main__":
    raise SystemExit(main())
